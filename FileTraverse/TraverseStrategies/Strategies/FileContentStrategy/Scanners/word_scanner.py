import warnings

from docx import Document

from FileTraverse.TraverseStrategies.Strategies.FileContentStrategy.abstract_file_scanner import AbstractFileScanner
from Utils.file_mime_type_enums import FileMimeTypeEnums


class WordScanner(AbstractFileScanner):
    def __init__(self):
        super().__init__()

    def scan(self, file_info):
        mime_type_enum = file_info["mime_type_enum"]
        match mime_type_enum:
            case FileMimeTypeEnums.doc:
                return self._scan_doc(file_info)
            case FileMimeTypeEnums.docx:
                return self._scan_docx(file_info)
        return False

    def support_scan(self, file_info):
        mime_type_enum = file_info["mime_type_enum"]
        match mime_type_enum:
            case FileMimeTypeEnums.docx:
                return True
            case FileMimeTypeEnums.doc:
                warnings.warn(f"不支持doc格式文档的读取，已跳过当前文件:{file_info["file_path"]}")
        return False

    def get_name(self):
        return "wordScanner"

    def _scan_doc(self, file_info):
        """
        暂不支持doc文件的读取
        """
        pass

    def _scan_docx(self, file_info):
        file_path = file_info["file_path"]
        # 加载Word文档
        doc = Document(file_path)
        # 遍历文档中的段落
        for para in doc.paragraphs:
            if self.content_match(para.text):
                return True

            # 遍历文档中的表格（如果需要）
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if self.content_match(cell.text):
                        return True

                    # 遍历文档中的标题（如果需要）
        for heading in doc.paragraphs:
            if heading.style.name.startswith('Heading'):
                if self.content_match(heading.text):
                    return True
